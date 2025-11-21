# -*- coding: utf-8 -*-
"""
密码验证函数测试用例
使用等价类划分和边界值分析法设计测试用例
"""

import pytest
from password_validator import isValidPassword

class TestPasswordValidator:
    """密码验证函数测试类"""
    
    def test_valid_passwords(self):
        """有效密码测试 - 等价类1: 符合所有要求的密码"""
        # 最短有效密码 (边界值: 6位)
        assert isValidPassword("abc123") == True
        
        # 最长有效密码 (边界值: 12位)
        assert isValidPassword("abcdef123456") == True
        
        # 中等长度有效密码
        assert isValidPassword("test1234") == True
        assert isValidPassword("abc123def") == True
        
        # 包含大小写字母和数字
        assert isValidPassword("Abc123") == True
        assert isValidPassword("Test123ABC") == True
    
    def test_invalid_length_too_short(self):
        """长度过短测试 - 等价类2: 长度小于6位"""
        # 边界值: 5位
        assert isValidPassword("ab123") == False
        
        # 更短的密码
        assert isValidPassword("a1") == False
        assert isValidPassword("123") == False
        assert isValidPassword("") == False  # 空字符串
    
    def test_invalid_length_too_long(self):
        """长度过长测试 - 等价类3: 长度大于12位"""
        # 边界值: 13位
        assert isValidPassword("abcdef1234567") == False
        
        # 更长的密码
        assert isValidPassword("abcdefghijk123456") == False
        assert isValidPassword("verylongpassword123") == False
    
    def test_no_digits(self):
        """不包含数字测试 - 等价类4: 只有字母"""
        # 边界长度但无数字
        assert isValidPassword("abcdef") == False  # 6位纯字母
        assert isValidPassword("abcdefghijkl") == False  # 12位纯字母
        
        # 中等长度纯字母
        assert isValidPassword("password") == False
        assert isValidPassword("ABCDEFGH") == False
        assert isValidPassword("AbCdEfGh") == False  # 混合大小写
    
    def test_no_letters(self):
        """不包含字母测试 - 等价类5: 只有数字"""
        # 边界长度但无字母
        assert isValidPassword("123456") == False  # 6位纯数字
        assert isValidPassword("123456789012") == False  # 12位纯数字
        
        # 中等长度纯数字
        assert isValidPassword("12345678") == False
        assert isValidPassword("987654321") == False
    
    def test_special_characters(self):
        """包含特殊字符测试 - 等价类6: 包含特殊字符"""
        # 有效密码包含特殊字符（题目未明确禁止）
        assert isValidPassword("abc123!") == True
        assert isValidPassword("test@123") == True
        assert isValidPassword("pass#word1") == True
        
        # 只有特殊字符和数字（无字母）
        assert isValidPassword("!@#123") == False
        
        # 只有特殊字符和字母（无数字）
        assert isValidPassword("abc!@#") == False
    
    def test_boundary_values(self):
        """边界值测试"""
        # 长度边界值测试
        assert isValidPassword("a1234") == False   # 5位 (长度-1)
        assert isValidPassword("a12345") == True   # 6位 (最小有效长度)
        assert isValidPassword("a123456789ab") == True  # 12位 (最大有效长度)
        assert isValidPassword("a123456789abc") == False # 13位 (长度+1)
        
        # 最小字符要求边界值
        assert isValidPassword("a12345") == True   # 最少字母数(1个)和数字数(1个)
        assert isValidPassword("ab1234") == True   # 多个字母
        assert isValidPassword("a12345") == True   # 多个数字
    
    def test_input_type_validation(self):
        """输入类型验证测试 - 等价类7: 非字符串输入"""
        # 非字符串输入
        assert isValidPassword(None) == False
        assert isValidPassword(123456) == False
        assert isValidPassword([]) == False
        assert isValidPassword({}) == False
        assert isValidPassword(True) == False
    
    def test_edge_cases(self):
        """边缘情况测试"""
        # 只在边界处满足条件
        assert isValidPassword("1a") == False      # 太短但有数字和字母
        assert isValidPassword("1abcdefghijk") == True   # 刚好12位
        assert isValidPassword("abcdefghijk1") == True   # 刚好12位，数字在末尾
        
        # 大小写混合
        assert isValidPassword("ABC123") == True
        assert isValidPassword("abc123") == True
        assert isValidPassword("AbC123") == True


def generate_test_report():
    """生成测试用例设计报告"""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('密码验证函数测试用例设计报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 功能规格
    doc.add_heading('1. 功能规格', level=1)
    
    p = doc.add_paragraph()
    p.add_run('函数名称：').bold = True
    p.add_run('isValidPassword(s: str) -> bool')
    
    p = doc.add_paragraph()
    p.add_run('功能描述：').bold = True
    p.add_run('验证输入的字符串是否为有效密码')
    
    p = doc.add_paragraph()
    p.add_run('输入参数：').bold = True
    p.add_run('s (str) - 待验证的密码字符串')
    
    p = doc.add_paragraph()
    p.add_run('返回值：').bold = True
    p.add_run('bool - 密码有效返回True，无效返回False')
    
    p = doc.add_paragraph()
    p.add_run('验证规则：').bold = True
    doc.add_paragraph('• 长度必须在6-12位之间', style='List Bullet')
    doc.add_paragraph('• 必须包含至少一个数字', style='List Bullet')
    doc.add_paragraph('• 必须包含至少一个字母（大小写均可）', style='List Bullet')
    
    # 等价类划分
    doc.add_heading('2. 等价类划分', level=1)
    
    equiv_table = doc.add_table(rows=1, cols=4)
    equiv_table.style = 'Table Grid'
    
    # 表头
    hdr_cells = equiv_table.rows[0].cells
    hdr_cells[0].text = '等价类编号'
    hdr_cells[1].text = '等价类描述'
    hdr_cells[2].text = '类型'
    hdr_cells[3].text = '测试数据示例'
    
    # 等价类数据
    equiv_data = [
        ('EC1', '长度6-12位，包含数字和字母', '有效等价类', '"abc123", "Test123ABC"'),
        ('EC2', '长度小于6位', '无效等价类', '"ab123", "a1"'),
        ('EC3', '长度大于12位', '无效等价类', '"abcdef1234567"'),
        ('EC4', '不包含数字', '无效等价类', '"abcdef", "password"'),
        ('EC5', '不包含字母', '无效等价类', '"123456", "987654"'),
        ('EC6', '包含特殊字符', '有效等价类', '"abc123!", "test@123"'),
        ('EC7', '非字符串输入', '无效等价类', 'None, 123456, []')
    ]
    
    for data in equiv_data:
        row_cells = equiv_table.add_row().cells
        for i, content in enumerate(data):
            row_cells[i].text = content
    
    # 边界值分析
    doc.add_heading('3. 边界值分析', level=1)
    
    boundary_table = doc.add_table(rows=1, cols=3)
    boundary_table.style = 'Table Grid'
    
    # 表头
    hdr_cells = boundary_table.rows[0].cells
    hdr_cells[0].text = '边界条件'
    hdr_cells[1].text = '边界值'
    hdr_cells[2].text = '测试用例'
    
    # 边界值数据
    boundary_data = [
        ('长度下边界', '6位', '"a12345" (有效), "a1234" (无效)'),
        ('长度上边界', '12位', '"a123456789ab" (有效), "a123456789abc" (无效)'),
        ('最少字符要求', '1个数字+1个字母', '"a12345" (最小有效组合)'),
        ('长度边界-1', '5位', '"ab123" (无效)'),
        ('长度边界+1', '13位', '"abcdef1234567" (无效)')
    ]
    
    for data in boundary_data:
        row_cells = boundary_table.add_row().cells
        for i, content in enumerate(data):
            row_cells[i].text = content
    
    # 测试用例汇总
    doc.add_heading('4. 测试用例汇总', level=1)
    
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = 'Table Grid'
    
    # 表头
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = '测试类别'
    hdr_cells[1].text = '用例数量'
    hdr_cells[2].text = '覆盖的等价类'
    hdr_cells[3].text = '预期结果'
    
    # 汇总数据
    summary_data = [
        ('有效密码测试', '6个', 'EC1, EC6', '全部返回True'),
        ('长度过短测试', '4个', 'EC2', '全部返回False'),
        ('长度过长测试', '3个', 'EC3', '全部返回False'),
        ('无数字测试', '5个', 'EC4', '全部返回False'),
        ('无字母测试', '4个', 'EC5', '全部返回False'),
        ('特殊字符测试', '6个', 'EC6', '部分True部分False'),
        ('边界值测试', '4个', '边界条件', '按边界规则'),
        ('类型验证测试', '5个', 'EC7', '全部返回False'),
        ('边缘情况测试', '6个', '多个等价类', '按规则判断')
    ]
    
    for data in summary_data:
        row_cells = summary_table.add_row().cells
        for i, content in enumerate(data):
            row_cells[i].text = content
    
    # 测试策略
    doc.add_heading('5. 测试策略', level=1)
    
    doc.add_paragraph('本测试用例设计采用了以下测试方法：')
    doc.add_paragraph('• 等价类划分：将输入域划分为7个等价类，确保每个类别都有代表性测试', style='List Bullet')
    doc.add_paragraph('• 边界值分析：重点测试长度边界（5,6,12,13位）和最小字符要求', style='List Bullet')
    doc.add_paragraph('• 组合测试：测试不同字符类型的组合情况', style='List Bullet')
    doc.add_paragraph('• 异常输入测试：验证非字符串输入的处理', style='List Bullet')
    
    doc.add_paragraph('总计设计了43个测试用例，覆盖了所有等价类和关键边界值，确保测试的完整性和有效性。')
    
    # 保存文档
    doc.save('密码验证函数测试用例设计报告.docx')
    print("测试用例设计报告已生成：密码验证函数测试用例设计报告.docx")


if __name__ == "__main__":
    # 运行测试
    pytest.main([__file__, "-v"])
    
    # 生成报告
    generate_test_report()
